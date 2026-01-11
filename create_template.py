from docxtpl import DocxTemplate
import os

def create_dummy_template():
    # Karena kita tidak bisa membuat file .docx valid dari nol hanya dengan text,
    # Kita akan menggunakan library docx (python-docx) jika ada, atau docxtpl untuk load basic.
    # Namun docxtpl butuh template dasar.
    # Alternatif: Kita buat file kosong atau minimal xml structure, tapi itu ribet.
    # Solusi terbaik: Gunakan python-docx untuk generate template awal.
    
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('RAPOR SISWA', 0)
        
        doc.add_paragraph('Nama Lengkap : {{ nama_lengkap }}')
        doc.add_paragraph('Kelas        : {{ Kelas }}')
        
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Mata Pelajaran'
        hdr_cells[1].text = 'Nilai'
        
        # Contoh baris dinamis manual (biasanya loop di jinja, tapi ini simple dulu)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Matematika'
        row_cells[1].text = '{{ matematika }}'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Bahasa Indonesia'
        row_cells[1].text = '{{ bahasa_indonesia }}'
        
        doc.add_paragraph('\nCatatan Wali Kelas:')
        doc.add_paragraph('{{ catatan_wali }}')
        
        output_path = os.path.join("templates", "rapor_template.docx")
        doc.save(output_path)
        print(f"Template dummy berhasil dibuat di: {output_path}")
        
    except ImportError:
        print("python-docx not found, skipping template creation. Please install python-docx first.")

if __name__ == "__main__":
    create_dummy_template()
